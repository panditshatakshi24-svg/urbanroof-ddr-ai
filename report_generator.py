from docx import Document

def generate_ddr(content, output_file):

    doc = Document()

    doc.add_heading(
        "Detailed Diagnostic Report",
        level=1
    )

    lines = str(content).split("\n")

    for line in lines:

        line = line.strip()

        if not line:
            continue

        if line in [
            "Property Issue Summary",
            "Area-wise Observations",
            "Probable Root Cause",
            "Severity Assessment",
            "Recommended Actions",
            "Additional Notes",
            "Missing or Unclear Information",
            "Conflicting Information"
        ]:
            doc.add_heading(line, level=2)

        else:
            doc.add_paragraph(line)

    doc.save(output_file)

    print(f"DDR saved to {output_file}")
